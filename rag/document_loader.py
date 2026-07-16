"""
Document loader – loads user-uploaded documents (PDF, DOCX, TXT, CSV)
into LangChain Document objects for RAG ingestion.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import List

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def load_text_file(path: Path) -> List[Document]:
    """Load a plain text file."""
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [Document(page_content=text, metadata={"source": str(path), "type": "text"})]
    except Exception as exc:
        logger.error("Failed to load text file %s: %s", path, exc)
        return []


def load_pdf(path: Path) -> List[Document]:
    """Load a PDF file using PyMuPDF (fitz) or pdfminer as fallback."""
    docs: List[Document] = []
    try:
        import fitz  # PyMuPDF

        pdf = fitz.open(str(path))
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text()
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": str(path), "page": page_num + 1, "type": "pdf"},
                ))
        pdf.close()
        return docs
    except ImportError:
        pass

    try:
        from pdfminer.high_level import extract_text

        text = extract_text(str(path))
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": str(path), "type": "pdf"},
            ))
        return docs
    except ImportError:
        logger.warning("Neither PyMuPDF nor pdfminer is installed. Cannot load PDF.")
        return []
    except Exception as exc:
        logger.error("PDF load failed for %s: %s", path, exc)
        return []


def load_docx(path: Path) -> List[Document]:
    """Load a DOCX file using python-docx."""
    try:
        import docx

        doc = docx.Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if text:
            return [Document(page_content=text, metadata={"source": str(path), "type": "docx"})]
        return []
    except ImportError:
        logger.warning("python-docx not installed. Cannot load DOCX.")
        return []
    except Exception as exc:
        logger.error("DOCX load failed for %s: %s", path, exc)
        return []


def load_csv(path: Path) -> List[Document]:
    """Load a CSV file as a text document."""
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [Document(page_content=text, metadata={"source": str(path), "type": "csv"})]
    except Exception as exc:
        logger.error("CSV load failed for %s: %s", path, exc)
        return []


def load_document(path: Path) -> List[Document]:
    """
    Load a document based on its extension.

    Supports: .txt, .md, .pdf, .docx, .csv

    Args:
        path: Path to the document file.

    Returns:
        List of LangChain Document objects.
    """
    path = Path(path)
    if not path.exists():
        logger.warning("Document not found: %s", path)
        return []

    ext = path.suffix.lower()
    loaders = {
        ".txt": load_text_file,
        ".md": load_text_file,
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".csv": load_csv,
    }
    loader_fn = loaders.get(ext)
    if loader_fn:
        docs = loader_fn(path)
        logger.info("Loaded %d document(s) from %s", len(docs), path)
        return docs
    else:
        logger.warning("Unsupported file type: %s", ext)
        return []


def load_documents_from_directory(directory: Path) -> List[Document]:
    """
    Load all supported documents from a directory.

    Args:
        directory: Directory path to scan.

    Returns:
        Combined list of all Document objects found.
    """
    directory = Path(directory)
    if not directory.is_dir():
        return []

    all_docs: List[Document] = []
    supported = {".txt", ".md", ".pdf", ".docx", ".csv"}
    for file_path in sorted(directory.iterdir()):
        if file_path.suffix.lower() in supported:
            all_docs.extend(load_document(file_path))

    logger.info("Loaded %d total document(s) from directory %s", len(all_docs), directory)
    return all_docs
